import os
import re
import sys

import docx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def add_hyperlink(par, label, url, size=None):
    part = par.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
    )
    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), r_id)
    run = par.add_run(label)
    run.font.color.rgb = RGBColor(0x0B, 0x4F, 0x9E)
    run.underline = True
    if size:
        run.font.size = Pt(size)
    par._p.remove(run._r)
    link.append(run._r)
    par._p.append(link)


INLINE = re.compile(
    r"(\*\*.+?\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\)|<https?://[^>]+>|\*[^*]+\*)"
)


def add_inline(par, text, size=None, code_font="Consolas"):
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = par.add_run(tok[2:-2])
            r.bold = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = code_font
        elif tok.startswith("[") and "](" in tok:
            label, url = tok[1:-1].split("](", 1)
            add_hyperlink(par, label, url, size=size)
            continue
        elif tok.startswith("<http"):
            url = tok[1:-1]
            add_hyperlink(par, url, url, size=size)
            continue
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            r = par.add_run(tok[1:-1])
            r.italic = True
        else:
            r = par.add_run(tok)
        if size:
            r.font.size = Pt(size)


def split_row(line):
    return [c.strip() for c in line.strip().strip("|").split("|")]


def is_sep(line):
    return bool(re.fullmatch(r"\|[\s:|-]+\|", line.strip()))


IMAGE_RE = re.compile(r"^!\[(?P<alt>[^\]]*)\]\((?P<path>[^)]+)\)$")


def _set_run_font(run, name=None, size=None, bold=None, italic=None):
    if name:
        run.font.name = name
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def apply_journal_styles(doc):
    """Science-style journal typography (single-column so wide tables stay readable)."""
    # Body text
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(10)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(6)

    # Lists should stay left-aligned
    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.space_after = Pt(3)

    # Headings
    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.font.bold = True
        style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        style.paragraph_format.line_spacing = 1.0
    doc.styles["Heading 1"].font.size = Pt(12)
    doc.styles["Heading 2"].font.size = Pt(10)
    doc.styles["Heading 2"].font.italic = True
    doc.styles["Heading 3"].font.size = Pt(10)
    doc.styles["Heading 3"].font.italic = True
    doc.styles["Heading 4"].font.size = Pt(10)

    # Title style for the first top-level heading
    title = doc.styles["Title"]
    title.font.name = "Times New Roman"
    title.font.size = Pt(14)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0, 0, 0)
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)

    # Margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)


def _shade_header_row(table, color="D9D9D9"):
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), color)
        tcPr.append(shd)


def convert(md_path, docx_path):
    base_dir = os.path.dirname(os.path.abspath(md_path))
    lines = open(md_path, encoding="utf-8").read().split("\n")
    doc = Document()
    apply_journal_styles(doc)

    i = 0
    n = len(lines)
    fig_num = 0
    while i < n:
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # horizontal rule
        if re.fullmatch(r"-{3,}|\*{3,}", stripped):
            i += 1
            continue

        # standalone image
        m_img = IMAGE_RE.match(stripped)
        if m_img:
            fig_num += 1
            alt = m_img.group("alt").strip()
            img_path = os.path.join(
                base_dir, m_img.group("path").replace("/", os.sep)
            )
            if os.path.exists(img_path):
                p = doc.add_picture(img_path, width=Inches(6.0))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(f"Figure {fig_num}. {alt}" if alt
                                  else f"Figure {fig_num}.")
                run.font.name = "Times New Roman"
                run.font.size = Pt(9)
                run.italic = True
            else:
                par = doc.add_paragraph()
                add_inline(par, f"[missing image: {m_img.group('path')}]")
            i += 1
            continue

        # fenced code block
        if stripped.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            par.paragraph_format.space_after = Pt(10)
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = par.add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        # table
        if stripped.startswith("|") and i + 1 < n and is_sep(lines[i + 1]):
            header = split_row(stripped)
            i += 2
            rows = []
            while i < n and lines[i].strip().startswith("|"):
                rows.append(split_row(lines[i].strip()))
                i += 1
            table = doc.add_table(rows=1, cols=len(header))
            table.style = "Table Grid"
            table.autofit = True
            for cell, text in zip(table.rows[0].cells, header):
                cell.paragraphs[0].text = ""
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_inline(cell.paragraphs[0], text, size=8)
                for r in cell.paragraphs[0].runs:
                    r.bold = True
            for row in rows:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.paragraphs[0].text = ""
                    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                    add_inline(cell.paragraphs[0], text, size=8)
            _shade_header_row(table)
            doc.add_paragraph()
            continue

        # headings
        m = re.match(r"(#{1,6})\s+(.*)", stripped)
        if m:
            level = len(m.group(1))
            par = doc.add_heading("", level=min(level, 4))
            add_inline(par, m.group(2))
            i += 1
            continue

        # blockquote
        if stripped.startswith(">"):
            buf = []
            while i < n and lines[i].strip().startswith(">"):
                buf.append(lines[i].strip().lstrip(">").strip())
                i += 1
            par = doc.add_paragraph()
            par.paragraph_format.left_indent = Pt(18)
            par.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            add_inline(par, " ".join(buf))
            for r in par.runs:
                r.italic = True
            continue

        # bullet list
        if re.match(r"[-*]\s+", stripped):
            par = doc.add_paragraph(style="List Bullet")
            add_inline(par, re.sub(r"^[-*]\s+", "", stripped))
            i += 1
            continue

        # numbered list
        if re.match(r"\d+\.\s+", stripped):
            par = doc.add_paragraph(style="List Number")
            add_inline(par, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        # paragraph (join consecutive lines)
        buf = [stripped]
        i += 1
        while i < n and lines[i].strip() and not re.match(
            r"(#{1,6}\s|[-*]\s|\d+\.\s|\||>|```|-{3,}$)", lines[i].strip()
        ):
            buf.append(lines[i].strip())
            i += 1
        par = doc.add_paragraph()
        add_inline(par, " ".join(buf))

    # Promote the first top-level heading to a centered title
    for p in doc.paragraphs:
        if p.style and p.style.name == "Heading 1":
            p.style = "Title"
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(14)
            break

    doc.save(docx_path)
    print("wrote", docx_path)


if __name__ == "__main__":
    convert(sys.argv[1], sys.argv[2])
